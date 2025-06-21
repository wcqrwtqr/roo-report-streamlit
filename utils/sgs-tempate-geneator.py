import streamlit as st
import os
from docx import Document

package_dir = os.path.dirname(os.path.abspath(__file__))

st.set_page_config(layout="wide")


def generate_sgs_page():
    """Read the data from kuster gauges over txt files and plot them.

    This function does not accept any parameters and excute
    a gauges data.
    """
    st.title("Geneate the SGS report_ 🌡")
    st.markdown(
        """
        Quickly and easily manipulate the data from __Kuster__ Down Hole Memory Gauges \
            further analysis or processing.
                """
    )
    source_data_bottom = st.file_uploader(
        label="Uplaod bottom gauge data to web page", type=["csv", "txt"],
        key="file_bottom_unique"
    )
    st.write("---")
    try:
        # Execute the program
        Gauges_data_kuster(source_data_bottom)
    except Exception as e:
        st.write("An error occured:" + str(e))



def replace_docx_variables(input_file, output_file, replacements):
    """Replace variables in DOCX file using python-docx."""
    document = Document(input_file)

    for paragraph in document.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    document.save(output_file)

st.title("Report Generator")

# Input fields using Streamlit widgets
wellname = st.text_input("Well Name", "Ru-524")
unit_no = st.selectbox("Unit Number", ["1", "2", "3"])  # Or st.number_input for integers
day = st.text_input("Day", "06")
activity = st.text_input("Activity", "Daily, Pre and Post job WSD")
activity_sgs = st.text_input("Activity SGS", "SGS")
activity_dr = st.text_input("Activity DR", "DR")
month = st.text_input("Month", "04")
year = st.text_input("Year", "2025")
monthfull = st.text_input("Month (Full)", "April")
dgs = st.text_input("DGS", "MDS")
tubing = st.text_input("Tubing", '4 1/2" 12.6# L-80')
sor = st.text_input("SOR", "34873")
welltype = st.text_input("Well Type", "Production")
fluid = st.text_input("Fluid", "Oil")
interval = st.text_input("Interval", "3388, 3370, 3364, 3350, 3342, 3332, 3264, 2989, 1989, 989, 489, 0mGL")
packer = st.text_input("Packer", "Packer")
field = st.text_input("Field", "South Rumaila")
min_res = st.text_input("Min Res", '3.44" no-go')
spv = st.text_input("SPV", "Malik Mohamed")


if st.button("Generate Report"):
    # Define input and output paths
    template_file = f"template{unit_no}.docx"  # Corrected path handling
    output_file = f"{wellname}_NEOS_SGS_Final-Report_{year}{month}{day}.docx"

    # Ensure the template file exists
    if not os.path.exists(template_file):
        st.error(f"Template file not found: {template_file}")
    else:

        # Create the replacements dictionary
        replacements = {
            "{{well_name}}": wellname,
            "{{sor}}": sor,
            "{{date}}": f"{day}-{month}-{year}",
            "{{type}}": "Static Gradient Survey",
            "{{field}}": field,
            "{{dgs}}": dgs,
            "{{fluid}}": fluid,
            "{{packer}}": packer,
            "{{tubing}}": tubing,
            "{{welltype}}": welltype,
            "{{interval}}": interval,
            "{{min-res}}": min_res,
            "{{spv}}": spv,
        }

        try:
            replace_docx_variables(template_file, output_file, replacements)
            st.success(f"Report generated successfully: {output_file}")

            # Offer download
            with open(output_file, "rb") as file:
                st.download_button(
                    label="Download Report",
                    data=file,
                    file_name=output_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


        except Exception as e:
            st.error(f"An error occurred: {e}") #Display the error message
