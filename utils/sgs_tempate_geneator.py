import streamlit as st
import os
from docx import Document
import pandas as pd

package_dir = os.path.dirname(os.path.abspath(__file__))

st.set_page_config(layout="wide")


def delete_file(filepath):
    """Deletes the file at the given filepath."""
    try:
        os.remove(filepath)
        print(f"File {filepath} deleted successfully.")  # Optional logging
    except Exception as e:
        print(f"Error deleting file {filepath}: {e}")  # Optional error logging
        st.error(f"Error deleting file {filepath}: {e}")


def generate_sgs_page():
    """Take the input data and generate a final report.  This is not final."""
    st.title("ROO SGS Report Generator")
    st.markdown(
        """
        Generate the SGS final report for ROO operation.
        Fill the fields to generate the final report quickly.
        """
    )

    # Define a dictionary to hold the default values and widget keys
    input_fields = {
        "wellname": {
            "label": "Well Name",
            "default": "Ru-524",
            "key": "wellname_input",
        },
        "welltype": {
            "label": "Well Type",
            "default": "Production",
            "key": "welltype_input",
        },
        "fluid": {"label": "Fluid", "default": "Oil", "key": "fluid_input"},
        "min_res": {
            "label": "Min Res",
            "default": '2.99"',
            "key": "min_res_input",
        },
        "spv": {"label": "SPV", "default": "Haider Jawad", "key": "spv_input"},
        "activity": {
            "label": "Activity",
            "default": "Daily, Pre and Post job WSD",
            "key": "activity_input",
        },
        "activity_sgs": {
            "label": "Activity SGS",
            "default": "SGS",
            "key": "activity_sgs_input",
        },
        "sgs-fgs": {
            "label": "SGS-FGS",
            "default": "Static Gradient Survey",
            "key": "sgs-fgs_input",
        },
        "packer": {"label": "Packer", "default": "Packer", "key": "packer_input"},
        "field": {"label": "Field", "default": "South Rumaila", "key": "field_input"},
        "sor": {"label": "SOR", "default": "34873", "key": "sor_input"},
        "day": {"label": "Day", "default": "06", "key": "day_input"},
        "casingsize": {
            "label": "Casing Size",
            "default": "7in 29# L-80",
            "key": "casingsize_input",
        },
        "month": {"label": "Month", "default": "04", "key": "month_input"},
        "monthfull": {
            "label": "Month (Full)",
            "default": "April",
            "key": "monthfull_input",
        },
        "year": {"label": "Year", "default": "2025", "key": "year_input"},
        "dgs": {"label": "DGS", "default": "MDS", "key": "dgs_input"},
        "tubing": {
            "label": "Tubing",
            "default": '3 1/2" 12.6# L-80',
            "key": "tubing_input",
        },
        "interval": {
            "label": "Interval",
            "default": "3388, 3342, 3332, 3264, 2989, 1989, 989, 489, 0mGL",
            "key": "interval_input",
        },
        "unit_no": {
            "label": "Unit Number",
            "options": ["1", "2", "3"],
            "default": "1",
            "key": "unit_no_input",
        },
    }

    # Selectbox for unit_no (has to be created differently)
    params = input_fields["unit_no"]
    if "unit_no" not in st.session_state:
        st.session_state["unit_no"] = params["default"]
    index = (
        params["options"].index(st.session_state["unit_no"])
        if st.session_state["unit_no"] in params["options"]
        else 0
    )

    col1, col2, col3 = st.columns(3)
    with col1:
        st.session_state["unit_no"] = st.selectbox(
            params["label"], params["options"], index=index, key=params["key"]
        )
        for field, params in input_fields.items():
            if field in ["wellname", "welltype", "fluid", "min_res", "spv"]:
                if field not in st.session_state:
                    st.session_state[field] = params["default"]
                st.session_state[field] = st.text_input(
                    params["label"], st.session_state[field], key=params["key"]
                )
    with col2:
        for field, params in input_fields.items():
            if field in [
                "activity",
                "field",
                "packer",
                "sor",
                "casingsize",
                "activity_sgs",
            ]:
                if field not in st.session_state:
                    st.session_state[field] = params["default"]
                st.session_state[field] = st.text_input(
                    params["label"], st.session_state[field], key=params["key"]
                )
    with col3:
        for field, params in input_fields.items():
            if field in ["day", "month", "year", "dgs", "tubing", "sgs-fgs"]:
                if field not in st.session_state:
                    st.session_state[field] = params["default"]
                st.session_state[field] = st.text_input(
                    params["label"], st.session_state[field], key=params["key"]
                )

    # Display interval input outside of the columns
    if "interval" not in st.session_state:
        st.session_state["interval"] = input_fields["interval"]["default"]
    st.session_state["interval"] = st.text_input(
        input_fields["interval"]["label"],
        st.session_state["interval"],
        key=input_fields["interval"]["key"],
    )

    if st.button("File Naming"):
        # This button will take the input and create the file names
        day = st.session_state["day"]
        activity = st.session_state["activity"]
        month = st.session_state["month"]
        year = st.session_state["year"]
        wellname = st.session_state["wellname"]
        unit_no = st.session_state["unit_no"]

        daily_report_name = f"Daily Report Name: {wellname}_NEOS_{
            activity
        }_Daily Report_{year}{month}{day}"
        pre_wsd_name = f"Pre Report Name: {wellname}_NEOS_{activity}_Pre Job WSD_{year}{
            month
        }{day}"
        post_wsd_name = f"Post Report Name: {wellname}_NEOS_{activity}_Post Job WSD_{
            year
        }{month}{day}"
        sgs_results_name = (
            f"SGS Report Name: {wellname}_NEOS_SGS Results_{year}{month}{day}"
        )
        if unit_no == "1":
            sgs_data_top_name = f"SGS Top Gauge data Name: {
                wellname
            }_NEOS_SGS Data_8701_top_{year}{month}{day}"
            sgs_data_bot_name = f"SGS Top Gauge data Name: {
                wellname
            }_NEOS_SGS Data_8700_bot_{year}{month}{day}"
        elif unit_no == "2":
            sgs_data_top_name = f"SGS Top Gauge data Name: {
                wellname
            }_NEOS_SGS Data_8698_top_{year}{month}{day}"
            sgs_data_bot_name = f"SGS Top Gauge data Name: {
                wellname
            }_NEOS_SGS Data_8699_bot_{year}{month}{day}"
        else:
            sgs_data_top_name = "Wrong input"
            sgs_data_bot_name = "Wrong input"

            # Create a list of dictionaries with two columns: Report Type and Report Name
        report_data = [
            {"Report Type": "Daily Report Name", "Report Name": daily_report_name},
            {"Report Type": "Pre Report Name", "Report Name": pre_wsd_name},
            {"Report Type": "Post Report Name", "Report Name": post_wsd_name},
            {"Report Type": "SGS Report Name", "Report Name": sgs_results_name},
            {
                "Report Type": "SGS Top Gauge Data Name",
                "Report Name": sgs_data_top_name,
            },
            {
                "Report Type": "SGS Bottom Gauge Data Name",
                "Report Name": sgs_data_bot_name,
            },
        ]

        # Convert to DataFrame
        df = pd.DataFrame(report_data)

        # Display as table
        st.table(df)

    if st.button("Generate Report"):
        # Get values from session state
        wellname = st.session_state["wellname"]
        sgs_fgs = st.session_state["sgs-fgs"]
        welltype = st.session_state["welltype"]
        fluid = st.session_state["fluid"]
        min_res = st.session_state["min_res"]
        spv = st.session_state["spv"]
        activity = st.session_state["activity"]
        # activity_sgs = st.session_state["activity_sgs"]
        packer = st.session_state["packer"]
        field = st.session_state["field"]
        sor = st.session_state["sor"]
        day = st.session_state["day"]
        month = st.session_state["month"]
        year = st.session_state["year"]
        dgs = st.session_state["dgs"]
        tubing = st.session_state["tubing"]
        interval = st.session_state["interval"]
        unit_no = st.session_state["unit_no"]
        casingsize = st.session_state["casingsize"]

        # Define input and output paths
        template_file = os.path.join("utils", f"template{unit_no}.docx")
        output_file = f"{
            wellname}_NEOS_SGS_Final-Report_{year}{month}{day}.docx"

        # Ensure the template file exists
        if not os.path.exists(template_file):
            st.error(f"Template file not found: {template_file}")
        else:
            # Create the replacements dictionary
            replacements = {
                "{{well_name}}": wellname,
                "{{sor}}": sor,
                "{{date}}": f"{day}-{month}-{year}",
                # "{{type}}": "Static Gradient Survey",
                "{{type}}": sgs_fgs,
                "{{field}}": field,
                "{{dgs}}": dgs,
                "{{fluid}}": fluid,
                "{{sgs-fgs}}": sgs_fgs,
                "{{packer}}": packer,
                "{{tubing}}": tubing,
                "{{welltype}}": welltype,
                "{{interval}}": interval,
                "{{min-res}}": min_res,
                "{{spv}}": spv,
                "{{casingsize}}": casingsize,
            }

            try:
                replace_docx_variables(
                    template_file, output_file, replacements)
                st.success(f"Report generated successfully: {output_file}")

                # Offer download
                with open(output_file, "rb") as file:
                    st.download_button(
                        label="Download Report",
                        data=file,
                        file_name=output_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        # Delete the file after download
                        on_click=delete_file,
                        kwargs={"filepath": output_file},
                    )

            except Exception as e:
                st.error(f"An error occurred: {e}")


def replace_docx_variables(input_file, output_file, replacements):
    """Replace variables in DOCX file, preserving formatting, using python-docx."""
    document = Document(input_file)

    def replace_text_in_paragraph(paragraph):
        """Replace text in a paragraph, preserving formatting of existing runs."""
        for key, value in replacements.items():
            if key in paragraph.text:
                new_value = str(value)  # Ensure value is a string
                inline = paragraph.runs
                # Loop added to work with runs (strings with same style)
                for i, run in enumerate(inline):
                    if key in run.text:
                        text = run.text.replace(key, new_value)
                        run.text = text

    def replace_text_in_cell(cell):
        """Replace text in a table cell, preserving formatting."""
        for paragraph in cell.paragraphs:
            # Apply paragraph-level replacement
            replace_text_in_paragraph(paragraph)

    # Replace in paragraphs in the main document
    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph)

    # Replace in tables in the main document
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell)

    # Replace in headers
    for section in document.sections:
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                replace_text_in_paragraph(
                    paragraph
                )  # Apply paragraph-level replacement
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Apply cell-level replacement
                        replace_text_in_cell(cell)

    # Replace in footers
    for section in document.sections:
        footer = section.footer
        if footer:
            for paragraph in footer.paragraphs:
                replace_text_in_paragraph(
                    paragraph
                )  # Apply paragraph-level replacement
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Apply cell-level replacement
                        replace_text_in_cell(cell)

    document.save(output_file)
